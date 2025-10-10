#!/usr/bin/env python3
# core/report_builders/word_report.py
import json, pathlib
from typing import Dict, Any, Tuple, List
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

Path = pathlib.Path

# --------------------------- helpers ---------------------------

def _abs_file_url(p: Path) -> str:
    ap = Path(p).resolve()
    return "file:///" + str(ap).replace("\\", "/")

def _add_hyperlink(paragraph, url: str, text: str, color="1155CC", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    if underline:
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    new_run.append(rPr)

    t = OxmlElement("w:t"); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def _resolve_screenshot(out_dir: Path, shot_val: str) -> Path:
    if not shot_val:
        return Path("")
    p = Path(str(shot_val).strip())
    if not p.is_absolute():
        p = (out_dir / p).resolve()
    return p if p.exists() else Path("")

def _load_json(p: Path, default):
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return default

def _add_kv_table(doc: Document, rows: List[Tuple[str, str]]):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)

def _add_table(doc: Document, headers: List[str], rows: List[List[str]], style: str = "Light List"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val if val is not None else "")

def _add_screenshot_block(doc: Document, out_dir: Path, shot_val: str):
    shot_path = _resolve_screenshot(out_dir, shot_val)
    if shot_path and shot_path.is_file():
        doc.add_paragraph("Screenshot:")
        try:
            doc.add_picture(str(shot_path), width=Inches(3.6))
        except Exception as e:
            doc.add_paragraph(f"(Could not embed image: {shot_path.name} — {e})")
        p = doc.add_paragraph()
        try:
            _add_hyperlink(p, _abs_file_url(shot_path), "Open full-size image")
        except Exception:
            p.add_run(str(shot_path))
    else:
        doc.add_paragraph("(No screenshot available)")

# --------------------------- main ---------------------------

def build_word(out_dir: Path, docx_path: Path, url: str):
    out_dir = Path(out_dir)

    axe = _load_json(out_dir / "axe_results.json", {})
    overall = axe.get("axe_issues", []) or []
    cands = _load_json(out_dir / "candidates.json", [])
    ai = _load_json(out_dir / "ai_verdicts.json", [])
    kb = _load_json(out_dir / "keyboard_probe.json", {})  # optional

    ai246_all = _load_json(out_dir / "ai" / "2_4_6" / "results.json", [])

    # For resolving screenshots if items didn’t carry one
    cand_by_selector: Dict[str, str] = {}
    for c in cands:
        sel = c.get("selector") or ""
        if sel and c.get("screenshot"):
            cand_by_selector.setdefault(sel, c["screenshot"])

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    # Title/meta
    doc.add_heading("Accessibility Audit Report", level=0)
    doc.add_paragraph(f"Target URL: {url}")
    doc.add_paragraph("Methodology: axe-core baseline, mechanical/hybrid detectors, keyboard probe, and semantic review (RAG + 2.4.6 AI).")

    # ---------------- Overview ----------------
    doc.add_heading("Overview", level=1)
    k_summary = kb.get("summary", {}) if isinstance(kb, dict) else {}
    _add_kv_table(doc, [
        ("Overall issues (mechanical + hybrid)", len(overall)),
        ("AI Candidates", len(cands)),
        ("AI Verdicts", len(ai)),
        ("AI 2.4.6 items", len(ai246_all)),
        ("Keyboard: Tab stops", k_summary.get("tab_stops","")),
        ("Keyboard: Unreachable", k_summary.get("unreachable","")),
        ("Keyboard: Activation checks", k_summary.get("activation_checks","")),
    ])

    # ---------------- NEW: WCAG 2.4.6 — Headings & Labels (AI) ----------------
    if ai246_all:
        doc.add_heading("WCAG 2.4.6 — Headings & Labels (AI)", level=1)
        verdicts = [str((x or {}).get("verdict","")).lower() for x in ai246_all]
        ai246_total = len(ai246_all)
        ai246_fail  = sum(1 for v in verdicts if v == "fail")
        ai246_pass  = sum(1 for v in verdicts if v == "pass")
        ai246_rev   = sum(1 for v in verdicts if v == "review")

        _add_kv_table(doc, [
            ("Total reviewed", ai246_total),
            ("Pass", ai246_pass),
            ("Fail", ai246_fail),
            ("Needs Review", ai246_rev),
        ])

        fails = [x for x in ai246_all if str(x.get("verdict","")).lower() == "fail"]
        if fails:
            doc.add_heading("Detailed Failures", level=2)
            for i, it in enumerate(fails, start=1):
                typ = it.get("type") or ""
                sel = it.get("selector") or ""
                reasons = it.get("reasons") or []
                sugg = it.get("suggestion") or ""
                visible = it.get("visibleText") or it.get("visibleLabel") or ""
                accname = it.get("accessibleName") or ""

                doc.add_heading(f"{i}. {typ.capitalize()} — {sel}", level=3)
                if visible:  doc.add_paragraph(f"Visible: {visible}")
                if accname:  doc.add_paragraph(f"Accessible name: {accname}")
                if reasons:
                    doc.add_paragraph("Reasons:")
                    for r in reasons:
                        doc.add_paragraph(f"• {r}")
                if sugg:
                    doc.add_paragraph(f"Suggestion: {sugg}")

                shot_val = it.get("screenshot") or cand_by_selector.get(sel) or ""
                _add_screenshot_block(doc, out_dir, shot_val)

        # Appendix table
        doc.add_page_break()
        doc.add_heading("Appendix — All 2.4.6 AI Judgements", level=1)
        headers = ["Type", "Selector", "Verdict", "Reasons", "Suggestion"]
        rows = []
        for it in ai246_all:
            rows.append([
                str(it.get("type","")),
                str(it.get("selector","")),
                str(it.get("verdict","")),
                "; ".join(it.get("reasons") or []),
                str(it.get("suggestion","") or ""),
            ])
        _add_table(doc, headers, rows)

    # ---------------- Overall Issues (mechanical & hybrid) ----------------
    doc.add_heading("Overall Issues (mechanical & hybrid)", level=1)
    if not overall:
        doc.add_paragraph("None recorded.")
    else:
        # Show FAIL items first, then REVIEW, then PASS (compact)
        def key_status(x):
            s = str(x.get("status","")).lower()
            return {"fail":0,"review":1,"pass":2}.get(s, 3)
        overall_sorted = sorted(overall, key=lambda x: (x.get("SC","") or x.get("sc",""), key_status(x), x.get("rule_id","") or x.get("axe_rule_id","")))
        for i, it in enumerate(overall_sorted, start=1):
            sc = it.get("SC") or it.get("sc") or ""
            rule = it.get("rule_id") or it.get("axe_rule_id") or ""
            status = (it.get("status") or it.get("verdict") or "").upper()
            sel = it.get("selector") or ""
            note = it.get("note") or it.get("failureSummary") or ""
            doc.add_heading(f"{i}. SC {sc} — {rule} [{status}]", level=2)
            if sel:   doc.add_paragraph(f"Selector: {sel}")
            if note:  doc.add_paragraph(f"Note: {note}")
            shot_val = it.get("screenshot") or cand_by_selector.get(sel) or ""
            _add_screenshot_block(doc, out_dir, shot_val)

    # ---------------- Findings (AI verdicts) ----------------
    doc.add_heading("Findings (AI verdicts)", level=1)
    if not ai:
        doc.add_paragraph("No AI verdicts were produced for this run.")
    else:
        for i, rec in enumerate(ai, start=1):
            topic = rec.get("topic") or rec.get("SC") or "Unmapped"
            rule  = rec.get("axe_rule_id") or ""
            sel   = rec.get("selector") or ""
            v     = rec.get("ai_verdict") or {}

            doc.add_heading(f"{i}. {topic} — {rule}", level=2)
            if sel:
                doc.add_paragraph(f"Selector: {sel}")
            verdict = v.get("verdict")
            conf    = v.get("confidence")
            if verdict is not None:
                doc.add_paragraph(f"Verdict: {verdict} (confidence {conf})")
            reason = v.get("reason")
            if reason:
                doc.add_paragraph(f"Reason: {reason}")
            if rec.get("axe_help_url"):
                doc.add_paragraph(f"Ref: {rec.get('axe_help_url')}")

            shot_val = rec.get("screenshot") or cand_by_selector.get(sel) or ""
            _add_screenshot_block(doc, out_dir, shot_val)

    # ---------------- Keyboard Probe ----------------
    doc.add_heading("Keyboard Probe", level=1)
    summ = kb.get("summary", {}) if kb else {}
    if summ:
        _add_kv_table(doc, [(k, summ[k]) for k in sorted(summ.keys())])

    # Detailed buckets (optional)
    if kb:
        def _render_kb_list(title: str, items: List[Dict[str,Any]], label_key: str = "selector"):
            if not items: return
            doc.add_heading(title, level=2)
            for i, it in enumerate(items, start=1):
                sel = it.get("selector") or ""
                role = it.get("role") or it.get("role_name_guess") or ""
                doc.add_heading(f"{i}. {sel}", level=3)
                if role: doc.add_paragraph(f"Role: {role}")
                shot_val = it.get("screenshot") or ""
                _add_screenshot_block(doc, out_dir, shot_val)

        _render_kb_list("Unreachable via Tab Order", kb.get("unreachable", []))
        _render_kb_list("Activation Checks (failed)", [x for x in kb.get("activations", []) if not (x.get("enter_ok") or x.get("space_ok"))])
        _render_kb_list('Elements with tabindex="-1"', kb.get("tabindex_neg1", []))

    doc.save(docx_path)
